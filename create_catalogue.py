#!/usr/bin/env python3
"""Generate a professional product catalogue .docx for Bê Tông Vải Cuộn."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(55, 65, 81)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.4


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set borders for a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{val[0]}" '
                f'w:space="0" w:color="{val[1]}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_styled_paragraph(doc_or_cell, text, size=11, bold=False, color=None, align=None, space_before=0, space_after=4):
    """Add a formatted paragraph."""
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Arial'
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_multi_run_paragraph(doc_or_cell, runs_data, align=None, space_before=0, space_after=4):
    """Add paragraph with multiple styled runs."""
    p = doc_or_cell.add_paragraph()
    for text, size, bold, color in runs_data:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = 'Arial'
        if color:
            run.font.color.rgb = RGBColor(*color)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ══════════════════════════════════════════════
# PAGE 1: COVER PAGE
# ══════════════════════════════════════════════

# Top accent bar
table_bar = doc.add_table(rows=1, cols=1)
table_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_bar.cell(0, 0)
set_cell_shading(cell, "1e40af")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(" ")
run.font.size = Pt(4)

# Add logo if available
logo_path = os.path.join(os.path.dirname(__file__), 'images', 'logo.jpg')
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run()
    run.add_picture(logo_path, width=Cm(4))

# Company name
add_styled_paragraph(doc, "CÔNG TY TNHH VẬN CHUYỂN QUỐC TẾ VIỆT NAM",
                     size=13, bold=True, color=(30, 64, 175), align='center', space_before=20)

# Decorative line
table_line = doc.add_table(rows=1, cols=1)
table_line.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_line.cell(0, 0)
set_cell_shading(cell, "f59e0b")
p = cell.paragraphs[0]
run = p.add_run(" ")
run.font.size = Pt(2)

# Main title
add_styled_paragraph(doc, "", size=6, space_after=0)
add_styled_paragraph(doc, "CATALOGUE SẢN PHẨM",
                     size=32, bold=True, color=(30, 64, 175), align='center', space_before=30)

add_styled_paragraph(doc, "BÊ TÔNG VẢI CUỘN",
                     size=28, bold=True, color=(5, 150, 105), align='center', space_before=8)

add_styled_paragraph(doc, "GEOSYNTHETIC CEMENTITIOUS COMPOSITE MAT (GCCM)",
                     size=12, bold=False, color=(107, 114, 128), align='center', space_before=10)

# Decorative divider
add_styled_paragraph(doc, "━━━━━━━━━━━━━━━━━━━━━━━━",
                     size=14, color=(229, 231, 235), align='center', space_before=20)

# Key highlights
highlights = [
    ("Tốc độ lắp đặt", "200m²/giờ"),
    ("Tuổi thọ công trình", "Lên đến 120 năm"),
    ("Tiêu chuẩn quốc tế", "ASTM D8364"),
    ("Tiết kiệm chi phí", "40-50% so với truyền thống"),
]

for label, value in highlights:
    add_multi_run_paragraph(doc, [
        ("●  ", 10, False, (245, 158, 11)),
        (f"{label}: ", 11, True, (55, 65, 81)),
        (value, 11, False, (5, 150, 105)),
    ], align='center', space_before=4, space_after=2)

# Contact info at bottom
add_styled_paragraph(doc, "", size=6, space_before=30)

table_contact = doc.add_table(rows=1, cols=1)
table_contact.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_contact.cell(0, 0)
set_cell_shading(cell, "f0f4ff")
add_styled_paragraph(cell, "Hotline: 0919 335 088  |  Email: betongvaicuon@gmail.com",
                     size=11, bold=True, color=(30, 64, 175), align='center', space_before=8, space_after=4)
add_styled_paragraph(cell, "Văn phòng: BT T02 - L07, Sol Lake Villa, Hà Đông, Hà Nội",
                     size=10, color=(107, 114, 128), align='center', space_before=0, space_after=8)

# Bottom accent bar
table_bar2 = doc.add_table(rows=1, cols=1)
table_bar2.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_bar2.cell(0, 0)
set_cell_shading(cell, "059669")
p = cell.paragraphs[0]
run = p.add_run(" ")
run.font.size = Pt(4)


# ══════════════════════════════════════════════
# PAGE 2: ABOUT & FEATURES
# ══════════════════════════════════════════════
doc.add_page_break()

# Header bar
table_h = doc.add_table(rows=1, cols=2)
table_h.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_l = table_h.cell(0, 0)
cell_r = table_h.cell(0, 1)
set_cell_shading(cell_l, "1e40af")
set_cell_shading(cell_r, "1e40af")
p = cell_l.paragraphs[0]
run = p.add_run("  BÊ TÔNG VẢI CUỘN")
run.font.size = Pt(9)
run.bold = True
run.font.color.rgb = RGBColor(255, 255, 255)
run.font.name = 'Arial'
p = cell_r.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("CATALOGUE SẢN PHẨM  ")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(200, 210, 240)
run.font.name = 'Arial'

add_styled_paragraph(doc, "GIỚI THIỆU SẢN PHẨM",
                     size=20, bold=True, color=(30, 64, 175), align='center', space_before=20)

# Amber underline
table_line2 = doc.add_table(rows=1, cols=1)
table_line2.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_line2.cell(0, 0)
set_cell_shading(cell, "f59e0b")
p = cell.paragraphs[0]
run = p.add_run(" ")
run.font.size = Pt(1)

add_styled_paragraph(doc, "", size=4)

intro_text = (
    "Bê tông vải cuộn (GCCM - Geosynthetic Cementitious Composite Mat) là vật liệu xây dựng "
    "thế hệ mới, kết hợp giữa vải địa kỹ thuật và xi măng khô trộn sẵn. Khi được phun nước, "
    "vật liệu sẽ hydrat hóa và đông cứng tạo thành lớp bê tông bền vững, chống thấm tuyệt đối."
)
p = add_styled_paragraph(doc, intro_text, size=11, color=(75, 85, 99), align='center', space_before=8)

add_styled_paragraph(doc, "", size=8)

# Features in 2-column table
add_styled_paragraph(doc, "ĐIỂM NỔI BẬT",
                     size=16, bold=True, color=(30, 64, 175), space_before=10, space_after=10)

features = [
    ("⚡ THI CÔNG NHANH CHÓNG", "Tốc độ lắp đặt 200m²/giờ, nhanh gấp 10 lần bê tông thông thường. Tiết kiệm thời gian và nhân công đáng kể."),
    ("💪 ĐỘ BỀN VƯỢT TRỘI", "Cường độ chịu nén tương đương MAC 600, tuổi thọ lên đến 120 năm theo chứng nhận BBA của Anh Quốc."),
    ("💧 CHỐNG THẤM TUYỆT ĐỐI", "Lớp nền PVC/LLDPE đảm bảo khả năng chống thấm hoàn hảo, ngăn ngừa xói mòn hiệu quả."),
    ("🎯 LINH HOẠT CAO", "Dễ dàng cắt, uốn theo địa hình. Phù hợp với mọi loại công trình từ đơn giản đến phức tạp."),
    ("🌱 THÂN THIỆN MÔI TRƯỜNG", "Không cần trạm trộn, giảm phát thải CO₂. Chỉ cần phun nước để kích hoạt quá trình đông kết."),
    ("💰 TIẾT KIỆM CHI PHÍ", "Giảm 40-50% chi phí so với phương pháp truyền thống nhờ thi công nhanh và ít tốn nhân công."),
]

feat_table = doc.add_table(rows=3, cols=2)
feat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (title, desc) in enumerate(features):
    row = i // 2
    col = i % 2
    cell = feat_table.cell(row, col)
    # Clear default paragraph
    cell.paragraphs[0].clear()
    set_cell_shading(cell, "f8fafc")
    set_cell_border(cell, top=("4", "e5e7eb"), bottom=("4", "e5e7eb"),
                    left=("4", "e5e7eb"), right=("4", "e5e7eb"))

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 64, 175)
    run.font.name = 'Arial'

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(107, 114, 128)
    run2.font.name = 'Arial'


# Applications section
add_styled_paragraph(doc, "", size=8)
add_styled_paragraph(doc, "ỨNG DỤNG ĐA DẠNG",
                     size=16, bold=True, color=(30, 64, 175), space_before=10, space_after=8)

applications = [
    "Lót kênh mương thủy lợi - Thay thế bê tông truyền thống",
    "Bảo vệ mái dốc, ta luy - Chống xói mòn, sạt lở hiệu quả",
    "Kè bờ sông, hồ chứa - Chịu sóng và dòng chảy cao",
    "Bọc ống dẫn, cống thoát - Tăng tuổi thọ hệ thống cống ngầm",
    "Công trình dầu khí, hóa chất - Đáp ứng tiêu chuẩn khắt khe",
    "Sửa chữa, nâng cấp công trình - Phục hồi nhanh công trình hư hỏng",
]

for app in applications:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("✓  ")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(5, 150, 105)
    run.bold = True
    run.font.name = 'Arial'
    parts = app.split(" - ")
    run2 = p.add_run(parts[0])
    run2.font.size = Pt(10)
    run2.bold = True
    run2.font.color.rgb = RGBColor(55, 65, 81)
    run2.font.name = 'Arial'
    if len(parts) > 1:
        run3 = p.add_run(f" - {parts[1]}")
        run3.font.size = Pt(10)
        run3.font.color.rgb = RGBColor(107, 114, 128)
        run3.font.name = 'Arial'


# ══════════════════════════════════════════════
# PAGE 3-4: PRODUCT CATALOGUE CARDS
# ══════════════════════════════════════════════
doc.add_page_break()

# Header bar
table_h2 = doc.add_table(rows=1, cols=2)
table_h2.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_l = table_h2.cell(0, 0)
cell_r = table_h2.cell(0, 1)
set_cell_shading(cell_l, "1e40af")
set_cell_shading(cell_r, "1e40af")
p = cell_l.paragraphs[0]
run = p.add_run("  BÊ TÔNG VẢI CUỘN")
run.font.size = Pt(9)
run.bold = True
run.font.color.rgb = RGBColor(255, 255, 255)
run.font.name = 'Arial'
p = cell_r.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("DANH MỤC SẢN PHẨM  ")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(200, 210, 240)
run.font.name = 'Arial'

add_styled_paragraph(doc, "DANH MỤC SẢN PHẨM",
                     size=20, bold=True, color=(30, 64, 175), align='center', space_before=20)

# Amber underline
table_line3 = doc.add_table(rows=1, cols=1)
table_line3.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_line3.cell(0, 0)
set_cell_shading(cell, "f59e0b")
p = cell.paragraphs[0]
run = p.add_run(" ")
run.font.size = Pt(1)

add_styled_paragraph(doc, "Dòng sản phẩm đa dạng về độ dày, đáp ứng mọi nhu cầu công trình.",
                     size=11, color=(107, 114, 128), align='center', space_before=8, space_after=12)

# Product data
products = [
    {
        "name": "Bê Tông Vải Cuộn 4mm",
        "code": "CC-4",
        "category": "Tiêu Chuẩn",
        "weight": "~6 kg/m²",
        "thickness": "4mm",
        "lifespan": "120 năm",
        "price": "100.000",
        "price_vat": "110.000",
        "color": "3b82f6",
        "apps": ["Lót kênh mương", "Chống thấm", "Mái dốc nhẹ"],
        "details": [
            "Cấu tạo: Lớp vải 3D Polyester + bê tông khô trộn sẵn + lớp chống thấm",
            "Phù hợp cho kênh mương nhỏ, mái dốc thoải, bọc ống cống",
            "Tốc độ thi công: 200m²/giờ",
            "Chỉ cần phun nước để kích hoạt đông kết xi măng",
        ],
    },
    {
        "name": "Bê Tông Vải Cuộn 6mm",
        "code": "CC-6",
        "category": "Tiêu Chuẩn",
        "badge": "PHỔ BIẾN",
        "weight": "~8 kg/m²",
        "thickness": "6mm",
        "lifespan": "120 năm",
        "price": "145.000",
        "price_vat": "159.500",
        "color": "059669",
        "apps": ["Kênh mương", "Bảo vệ mái dốc", "Kè bờ"],
        "details": [
            "Lựa chọn phổ biến nhất cho các công trình vừa và nhỏ",
            "Phù hợp: lót kênh, mương tưới tiêu, bảo vệ mái taluy",
            "Cường độ chịu nén tương đương MAC 600",
            "Khả năng chống thấm hoàn hảo với lớp nền PVC/LLDPE",
        ],
    },
    {
        "name": "Bê Tông Vải Cuộn 8mm",
        "code": "CC-8",
        "category": "Tiêu Chuẩn",
        "weight": "~10 kg/m²",
        "thickness": "8mm",
        "lifespan": "120 năm",
        "price": "165.000",
        "price_vat": "181.500",
        "color": "b45309",
        "apps": ["Ta luy đường", "Kè bờ sông", "Hồ chứa"],
        "details": [
            "Độ dày trung bình, cân bằng giữa hiệu quả và chi phí",
            "Lý tưởng cho kè bờ sông, hồ chứa nước, ta luy đường bộ",
            "Chịu được sóng và dòng chảy trung bình",
            "Giảm 40-50% chi phí so với bê tông truyền thống",
        ],
    },
    {
        "name": "Bê Tông Vải Cuộn 10mm",
        "code": "CC-10",
        "category": "Cao Cấp",
        "badge": "BÁN CHẠY",
        "weight": "~12 kg/m²",
        "thickness": "10mm",
        "lifespan": "120 năm",
        "price": "185.000",
        "price_vat": "203.500",
        "color": "4338ca",
        "apps": ["Công trình lớn", "Dầu khí", "Hóa chất"],
        "details": [
            "Độ bền cao, phù hợp công trình hạ tầng quan trọng",
            "Ứng dụng: kè bờ sông lớn, hồ chứa, công trình dầu khí",
            "Chịu được tải trọng lớn và điều kiện khắc nghiệt",
            "Tuổi thọ lên đến 120 năm (chứng nhận BBA)",
        ],
    },
    {
        "name": "Bê Tông Vải Cuộn 12mm",
        "code": "CC-12",
        "category": "Cao Cấp",
        "weight": "~15 kg/m²",
        "thickness": "12mm",
        "lifespan": "120 năm",
        "price": "205.000",
        "price_vat": "225.500",
        "color": "be185d",
        "apps": ["Hạ tầng lớn", "Chống sạt lở", "Ven biển"],
        "details": [
            "Dành cho các dự án hạ tầng quy mô lớn, yêu cầu cao",
            "Ứng dụng: kè biển, đê điều, công trình ven sông lớn",
            "Chống sạt lở, chịu sóng biển và dòng chảy mạnh",
            "Thi công được cả trong điều kiện thời tiết khắc nghiệt",
        ],
    },
    {
        "name": "Bê Tông Vải Cuộn 15mm",
        "code": "CC-15",
        "category": "Cao Cấp",
        "weight": "~18 kg/m²",
        "thickness": "15mm",
        "lifespan": "120 năm",
        "price": "265.000",
        "price_vat": "291.500",
        "color": "b91c1c",
        "apps": ["Đê biển", "Quốc phòng", "Hầm mỏ"],
        "details": [
            "Độ dày tối đa, bảo vệ công trình ở mức cao nhất",
            "Dành cho: đê biển, công trình quốc phòng, hầm mỏ",
            "Chịu được va đập, tải trọng siêu lớn",
            "Phù hợp cho vùng có điều kiện địa chất phức tạp",
        ],
    },
    {
        "name": "Bê Tông Vải Cuộn 10mm Premium",
        "code": "CC-10P",
        "category": "Premium",
        "badge": "PREMIUM",
        "weight": "~15 kg/m²",
        "thickness": "10mm",
        "lifespan": "120+ năm",
        "price": "290.000",
        "price_vat": "319.000",
        "color": "7c3aed",
        "apps": ["3D Polyester", "PVC chống thấm", "Sợi gia cường"],
        "details": [
            "Lớp 1: Lưới 3D Polyester - Tạo khung cấu trúc vững chắc",
            "Lớp 2: Sợi gia cường - Tăng cường khả năng chịu kéo",
            "Lớp 3: Bê tông khô trộn sẵn - Đông kết khi tiếp xúc nước",
            "Lớp 4: Lớp chống thấm PVC - Ngăn thấm tuyệt đối",
        ],
    },
]


# Render product cards (2 per row in a table)
for idx in range(0, len(products), 2):
    if idx > 0 and idx % 4 == 0:
        doc.add_page_break()
        # Re-add header bar
        table_h3 = doc.add_table(rows=1, cols=2)
        table_h3.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_l = table_h3.cell(0, 0)
        cell_r = table_h3.cell(0, 1)
        set_cell_shading(cell_l, "1e40af")
        set_cell_shading(cell_r, "1e40af")
        p = cell_l.paragraphs[0]
        run = p.add_run("  BÊ TÔNG VẢI CUỘN")
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Arial'
        p = cell_r.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("DANH MỤC SẢN PHẨM  ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(200, 210, 240)
        run.font.name = 'Arial'
        add_styled_paragraph(doc, "", size=6, space_after=6)

    batch = products[idx:idx + 2]
    num_cols = len(batch)

    card_table = doc.add_table(rows=1, cols=num_cols)
    card_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_idx, prod in enumerate(batch):
        cell = card_table.cell(0, col_idx)
        set_cell_shading(cell, "f8fafc")
        set_cell_border(cell,
                        top=("6", prod["color"]),
                        bottom=("4", "e5e7eb"),
                        left=("4", "e5e7eb"),
                        right=("4", "e5e7eb"))

        # Badge + Code
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        badge_text = prod.get("badge", "")
        run = p.add_run(f"  {prod['code']}")
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(107, 114, 128)
        run.font.name = 'Arial'
        if badge_text:
            run2 = p.add_run(f"   ★ {badge_text}")
            run2.font.size = Pt(8)
            run2.bold = True
            c = prod["color"]
            run2.font.color.rgb = RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
            run2.font.name = 'Arial'

        # Product name
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(2)
        run = p2.add_run(prod["name"])
        run.font.size = Pt(13)
        run.bold = True
        c = prod["color"]
        run.font.color.rgb = RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
        run.font.name = 'Arial'

        # Category
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(6)
        run = p3.add_run(f"Dòng {prod['category']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(156, 163, 175)
        run.font.name = 'Arial'

        # Specs mini-table
        spec_table = cell.add_table(rows=2, cols=3)
        specs = [
            ("Trọng lượng", prod["weight"]),
            ("Độ dày", prod["thickness"]),
            ("Tuổi thọ", prod["lifespan"]),
        ]
        for si, (label, val) in enumerate(specs):
            sc = spec_table.cell(0, si)
            sc.paragraphs[0].clear()
            set_cell_shading(sc, "eef2ff")
            p = sc.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 64, 175)
            run.font.name = 'Arial'

            sc2 = spec_table.cell(1, si)
            sc2.paragraphs[0].clear()
            set_cell_shading(sc2, "eef2ff")
            p = sc2.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(label)
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(156, 163, 175)
            run.font.name = 'Arial'

        # Application tags
        p_tags = cell.add_paragraph()
        p_tags.paragraph_format.space_before = Pt(8)
        p_tags.paragraph_format.space_after = Pt(4)
        for app in prod["apps"]:
            run = p_tags.add_run(f" {app} ")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(30, 64, 175)
            run.font.name = 'Arial'
            run = p_tags.add_run("  ")
            run.font.size = Pt(8)

        # Details
        for detail in prod["details"]:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run("● ")
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(5, 150, 105)
            run.font.name = 'Arial'
            run2 = p.add_run(detail)
            run2.font.size = Pt(8)
            run2.font.color.rgb = RGBColor(107, 114, 128)
            run2.font.name = 'Arial'

        # Price
        p_price = cell.add_paragraph()
        p_price.paragraph_format.space_before = Pt(10)
        p_price.paragraph_format.space_after = Pt(8)
        run = p_price.add_run(f"{prod['price']} VNĐ/m²")
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(5, 150, 105)
        run.font.name = 'Arial'
        run2 = p_price.add_run(f"  (VAT: {prod['price_vat']}đ)")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(156, 163, 175)
        run2.font.name = 'Arial'

    add_styled_paragraph(doc, "", size=6, space_after=4)


# ══════════════════════════════════════════════
# PAGE 5: PRICE TABLE + NOTES
# ══════════════════════════════════════════════
doc.add_page_break()

# Header bar
table_h4 = doc.add_table(rows=1, cols=2)
table_h4.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_l = table_h4.cell(0, 0)
cell_r = table_h4.cell(0, 1)
set_cell_shading(cell_l, "1e40af")
set_cell_shading(cell_r, "1e40af")
p = cell_l.paragraphs[0]
run = p.add_run("  BÊ TÔNG VẢI CUỘN")
run.font.size = Pt(9)
run.bold = True
run.font.color.rgb = RGBColor(255, 255, 255)
run.font.name = 'Arial'
p = cell_r.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("BẢNG BÁO GIÁ  ")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(200, 210, 240)
run.font.name = 'Arial'

add_styled_paragraph(doc, "BẢNG BÁO GIÁ SẢN PHẨM",
                     size=20, bold=True, color=(30, 64, 175), align='center', space_before=20)

table_line4 = doc.add_table(rows=1, cols=1)
table_line4.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_line4.cell(0, 0)
set_cell_shading(cell, "f59e0b")
p = cell.paragraphs[0]
run = p.add_run(" ")
run.font.size = Pt(1)

add_styled_paragraph(doc, "", size=8)

# Price table
price_table = doc.add_table(rows=8, cols=5)
price_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Sản phẩm", "Xuất xứ", "ĐVT", "Đơn giá (VNĐ)", "Giá có VAT (VNĐ)"]
price_data = [
    ["Bê tông vải cuộn 4mm\n(~6kg/m²)", "Trung Quốc", "m²", "100.000", "110.000"],
    ["Bê tông vải cuộn 6mm\n(~8kg/m²)", "Trung Quốc", "m²", "145.000", "159.500"],
    ["Bê tông vải cuộn 8mm\n(~10kg/m²)", "Trung Quốc", "m²", "165.000", "181.500"],
    ["Bê tông vải cuộn 10mm\n(~12kg/m²)", "Trung Quốc", "m²", "185.000", "203.500"],
    ["Bê tông vải cuộn 12mm\n(~15kg/m²)", "Trung Quốc", "m²", "205.000", "225.500"],
    ["Bê tông vải cuộn 15mm\n(~18kg/m²)", "Trung Quốc", "m²", "265.000", "291.500"],
    ["Bê tông vải cuộn 10mm\nPremium (~15kg/m²)", "Trung Quốc", "m²", "290.000", "319.000"],
]

# Header row
for ci, header in enumerate(headers):
    cell = price_table.cell(0, ci)
    set_cell_shading(cell, "1e40af")
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(header)
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = 'Arial'

# Data rows
for ri, row_data in enumerate(price_data):
    bg = "f8fafc" if ri % 2 == 0 else "ffffff"
    for ci, val in enumerate(row_data):
        cell = price_table.cell(ri + 1, ci)
        set_cell_shading(cell, bg)
        set_cell_border(cell, bottom=("2", "e5e7eb"))
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        if ci == 0:
            lines = val.split("\n")
            run = p.add_run(lines[0])
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 64, 175)
            run.font.name = 'Arial'
            if len(lines) > 1:
                run2 = p.add_run(f"\n{lines[1]}")
                run2.font.size = Pt(8)
                run2.font.color.rgb = RGBColor(156, 163, 175)
                run2.font.name = 'Arial'
        elif ci == 3:
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(5, 150, 105)
            run.font.name = 'Arial'
        elif ci == 4:
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(220, 38, 38)
            run.font.name = 'Arial'
        else:
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(55, 65, 81)
            run.font.name = 'Arial'

# Notes section
add_styled_paragraph(doc, "", size=8)

note_table = doc.add_table(rows=1, cols=1)
note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = note_table.cell(0, 0)
set_cell_shading(cell, "fffbeb")
set_cell_border(cell, left=("12", "f59e0b"), top=("4", "fde68a"),
                bottom=("4", "fde68a"), right=("4", "fde68a"))

p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("⚠  LƯU Ý QUAN TRỌNG")
run.font.size = Pt(12)
run.bold = True
run.font.color.rgb = RGBColor(146, 64, 14)
run.font.name = 'Arial'

notes = [
    "Giá có thể thay đổi theo thời điểm cụ thể và địa điểm giao hàng",
    "Giá trên là giá nhận hàng tại Hà Nội",
    "Thuế suất VAT: 10%",
    "Thanh toán: Đặt cọc 80%, thanh toán phần còn lại sau khi bàn giao đầy đủ chứng từ",
    "Thời gian giao hàng: 15-30 ngày kể từ ngày chuyển tiền tạm ứng",
    "Vui lòng liên hệ để nhận báo giá chính xác nhất",
]

for note in notes:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("•  ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(245, 158, 11)
    run.bold = True
    run.font.name = 'Arial'
    run2 = p.add_run(note)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(120, 53, 15)
    run2.font.name = 'Arial'

# Green CTA box
add_styled_paragraph(doc, "", size=6)

cta_table = doc.add_table(rows=1, cols=1)
cta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = cta_table.cell(0, 0)
set_cell_shading(cell, "059669")

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
run = p.add_run("Liên hệ ngay: 0919 335 088 để được tư vấn và báo giá tốt nhất!")
run.font.size = Pt(13)
run.bold = True
run.font.color.rgb = RGBColor(255, 255, 255)
run.font.name = 'Arial'


# ══════════════════════════════════════════════
# PAGE 6: TECHNICAL SPECS & STANDARDS
# ══════════════════════════════════════════════
doc.add_page_break()

# Header bar
table_h5 = doc.add_table(rows=1, cols=2)
table_h5.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_l = table_h5.cell(0, 0)
cell_r = table_h5.cell(0, 1)
set_cell_shading(cell_l, "1e40af")
set_cell_shading(cell_r, "1e40af")
p = cell_l.paragraphs[0]
run = p.add_run("  BÊ TÔNG VẢI CUỘN")
run.font.size = Pt(9)
run.bold = True
run.font.color.rgb = RGBColor(255, 255, 255)
run.font.name = 'Arial'
p = cell_r.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("THÔNG SỐ KỸ THUẬT  ")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(200, 210, 240)
run.font.name = 'Arial'

add_styled_paragraph(doc, "THÔNG SỐ KỸ THUẬT",
                     size=20, bold=True, color=(30, 64, 175), align='center', space_before=20)

table_line5 = doc.add_table(rows=1, cols=1)
table_line5.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_line5.cell(0, 0)
set_cell_shading(cell, "f59e0b")
p = cell.paragraphs[0]
run = p.add_run(" ")
run.font.size = Pt(1)

add_styled_paragraph(doc, "", size=8)

# Technical comparison table
tech_table = doc.add_table(rows=9, cols=5)
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

tech_headers = ["Thông số", "CC-4 / CC-6", "CC-8 / CC-10", "CC-12 / CC-15", "CC-10P"]
tech_data = [
    ["Độ dày (mm)", "4 / 6", "8 / 10", "12 / 15", "10"],
    ["Trọng lượng (kg/m²)", "~6 / ~8", "~10 / ~12", "~15 / ~18", "~15"],
    ["Cường độ chịu nén", "MAC 600", "MAC 600", "MAC 600", "MAC 600+"],
    ["Lớp chống thấm", "PVC/LLDPE", "PVC/LLDPE", "PVC/LLDPE", "PVC cao cấp"],
    ["Số lớp cấu tạo", "3 lớp", "3 lớp", "3 lớp", "4 lớp"],
    ["Tuổi thọ (năm)", "≥ 120", "≥ 120", "≥ 120", "≥ 120"],
    ["Tiêu chuẩn", "ASTM D8364", "ASTM D8364", "ASTM D8364", "ASTM D8364"],
    ["Phân loại ASTM", "CCT1™ Type I", "CCT2™ Type II", "CCT3™ Type III", "CCT2™ Type II+"],
]

for ci, h in enumerate(tech_headers):
    cell = tech_table.cell(0, ci)
    set_cell_shading(cell, "1e40af")
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(h)
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = 'Arial'

for ri, row in enumerate(tech_data):
    bg = "f8fafc" if ri % 2 == 0 else "ffffff"
    for ci, val in enumerate(row):
        cell = tech_table.cell(ri + 1, ci)
        set_cell_shading(cell, bg)
        set_cell_border(cell, bottom=("2", "e5e7eb"))
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        if ci == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(55, 65, 81)
        else:
            run.font.color.rgb = RGBColor(75, 85, 99)

# Standards section
add_styled_paragraph(doc, "", size=10)
add_styled_paragraph(doc, "CHỨNG NHẬN & TIÊU CHUẨN",
                     size=16, bold=True, color=(30, 64, 175), space_before=10, space_after=10)

standards = [
    ("ASTM D8364", "Tiêu chuẩn quốc tế về vật liệu GCCM (Geosynthetic Cementitious Composite Mat)"),
    ("BBA (Anh Quốc)", "Chứng nhận tuổi thọ sản phẩm lên đến 120 năm"),
    ("CCT1™ - Type I", "Phân loại sản phẩm dòng tiêu chuẩn (4-6mm)"),
    ("CCT2™ - Type II", "Phân loại sản phẩm dòng trung bình (8-10mm)"),
    ("CCT3™ - Type III", "Phân loại sản phẩm dòng cao cấp (12-15mm)"),
]

for std_name, std_desc in standards:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("✅  ")
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run2 = p.add_run(std_name)
    run2.font.size = Pt(10)
    run2.bold = True
    run2.font.color.rgb = RGBColor(30, 64, 175)
    run2.font.name = 'Arial'
    run3 = p.add_run(f"  —  {std_desc}")
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(107, 114, 128)
    run3.font.name = 'Arial'


# ══════════════════════════════════════════════
# PAGE 7: CONTACT
# ══════════════════════════════════════════════
doc.add_page_break()

# Header bar
table_h6 = doc.add_table(rows=1, cols=2)
table_h6.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_l = table_h6.cell(0, 0)
cell_r = table_h6.cell(0, 1)
set_cell_shading(cell_l, "1e40af")
set_cell_shading(cell_r, "1e40af")
p = cell_l.paragraphs[0]
run = p.add_run("  BÊ TÔNG VẢI CUỘN")
run.font.size = Pt(9)
run.bold = True
run.font.color.rgb = RGBColor(255, 255, 255)
run.font.name = 'Arial'
p = cell_r.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("LIÊN HỆ  ")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(200, 210, 240)
run.font.name = 'Arial'

add_styled_paragraph(doc, "", size=30)

add_styled_paragraph(doc, "LIÊN HỆ VỚI CHÚNG TÔI",
                     size=24, bold=True, color=(30, 64, 175), align='center', space_before=30)

table_line6 = doc.add_table(rows=1, cols=1)
table_line6.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_line6.cell(0, 0)
set_cell_shading(cell, "f59e0b")
p = cell.paragraphs[0]
run = p.add_run(" ")
run.font.size = Pt(1)

add_styled_paragraph(doc, "Đội ngũ kỹ thuật sẵn sàng tư vấn giải pháp phù hợp nhất cho công trình của bạn.",
                     size=12, color=(107, 114, 128), align='center', space_before=15, space_after=25)

# Contact card
contact_table = doc.add_table(rows=1, cols=1)
contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = contact_table.cell(0, 0)
set_cell_shading(cell, "f0f4ff")
set_cell_border(cell, top=("6", "1e40af"), bottom=("4", "e0e7ff"),
                left=("4", "e0e7ff"), right=("4", "e0e7ff"))

# Company name
cell.paragraphs[0].clear()
add_styled_paragraph(cell, "CÔNG TY TNHH VẬN CHUYỂN QUỐC TẾ VIỆT NAM",
                     size=14, bold=True, color=(30, 64, 175), align='center', space_before=20, space_after=15)

contact_items = [
    ("📞  Hotline:", "0919 335 088"),
    ("✉️   Email:", "betongvaicuon@gmail.com"),
    ("🌐  Website:", "betongvaicuon.com"),
    ("📍  Văn phòng:", "BT T02 - L07, Sol Lake Villa, Hà Đông, Hà Nội, Việt Nam"),
    ("🏭  Nhà máy:", "Quảng Đông, Trung Quốc"),
    ("⏰  Giờ làm việc:", "Thứ 2 - Thứ 7: 8:00 - 17:30"),
]

for icon_label, value in contact_items:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(icon_label + "  ")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(55, 65, 81)
    run.font.name = 'Arial'
    run2 = p.add_run(value)
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(30, 64, 175)
    run2.font.name = 'Arial'

add_styled_paragraph(cell, "", size=8)

# Zalo CTA
add_styled_paragraph(doc, "", size=15)

zalo_table = doc.add_table(rows=1, cols=1)
zalo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = zalo_table.cell(0, 0)
set_cell_shading(cell, "059669")

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(15)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("GỌI NGAY: 0919 335 088")
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(255, 255, 255)
run.font.name = 'Arial'

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(15)
run = p2.add_run("Tư vấn miễn phí  •  Báo giá nhanh  •  Hỗ trợ kỹ thuật 24/7")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(209, 250, 229)
run.font.name = 'Arial'

# Footer
add_styled_paragraph(doc, "", size=20)
add_styled_paragraph(doc, "© 2025 Bê Tông Vải Cuộn Việt Nam. Bản quyền thuộc về chúng tôi.",
                     size=9, color=(156, 163, 175), align='center', space_before=20)


# ══════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), "Catalogue_BeTongVaiCuon_2025.docx")
doc.save(output_path)
print(f"✅ Catalogue saved to: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024:.1f} KB")
print(f"   Pages: ~7 pages")
